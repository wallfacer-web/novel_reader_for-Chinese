#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
英文小说阅读辅助软件 (English Novel Reading Assistant)
一个基于AI的英文小说阅读辅助工具，帮助中文用户更好地理解和学习英文小说。

Author: Toby LUO@ZHKU (903098625@qq.com)
Copyright (c) 2024 Toby LUO@ZHKU (903098625@qq.com)
License: MIT License

GitHub: https://github.com/wallfacer-web/novel_reader_for-Chinese
"""

import gradio as gr
import requests
import json
import re
from typing import List, Dict, Any, Tuple
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import logging
import math
from collections import Counter
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
import sqlite3
from datetime import datetime

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 下载必要的NLTK数据
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

class VocabularyDatabase:
    """词汇数据库管理"""
    
    def __init__(self, db_path: str = "vocabulary.db"):
        self.db_path = db_path
        self.init_database()
    
    def init_database(self):
        """初始化数据库"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # 创建词汇表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS vocabulary (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                word TEXT UNIQUE,
                definition TEXT,
                word_family TEXT,
                frequency_level INTEGER,
                learned_count INTEGER DEFAULT 0,
                first_seen DATE,
                last_reviewed DATE
            )
        ''')
        
        # 创建学习记录表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS learning_progress (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_date DATE,
                words_learned INTEGER,
                paragraphs_processed INTEGER,
                reading_time INTEGER
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def add_word(self, word: str, definition: str, word_family: str = "", frequency_level: int = 5):
        """添加单词到数据库"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                INSERT OR REPLACE INTO vocabulary 
                (word, definition, word_family, frequency_level, first_seen)
                VALUES (?, ?, ?, ?, ?)
            ''', (word.lower(), definition, word_family, frequency_level, datetime.now().date()))
            conn.commit()
        except Exception as e:
            logger.error(f"添加单词失败: {e}")
        finally:
            conn.close()
    
    def get_learned_words(self) -> List[str]:
        """获取已学习的单词列表"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('SELECT word FROM vocabulary WHERE learned_count > 0')
        words = [row[0] for row in cursor.fetchall()]
        
        conn.close()
        return words

class TextDifficultyAnalyzer:
    """文本难度分析器"""
    
    def __init__(self):
        # 基础常用词汇表（模拟前3000个最常用英语单词）
        self.common_words = self._load_basic_words()
    
    def _load_basic_words(self) -> set:
        """加载基础词汇表"""
        basic_words = [
            'the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have',
            'i', 'it', 'for', 'not', 'on', 'with', 'he', 'as', 'you',
            'do', 'at', 'this', 'but', 'his', 'by', 'from', 'they',
            'she', 'or', 'an', 'will', 'my', 'one', 'all', 'would',
            'there', 'their', 'what', 'so', 'up', 'out', 'if', 'about',
            'who', 'get', 'which', 'go', 'me', 'when', 'make', 'can',
            'like', 'time', 'no', 'just', 'him', 'know', 'take', 'people',
            'into', 'year', 'your', 'good', 'some', 'could', 'them', 'see',
            'other', 'than', 'then', 'now', 'look', 'only', 'come', 'its',
            'over', 'think', 'also', 'back', 'after', 'use', 'two', 'how',
            'our', 'work', 'first', 'well', 'way', 'even', 'new', 'want',
            'because', 'any', 'these', 'give', 'day', 'most', 'us', 'was',
            'been', 'said', 'each', 'which', 'she', 'do', 'how', 'their',
            'if', 'will', 'up', 'other', 'about', 'out', 'many', 'then',
            'them', 'these', 'so', 'some', 'her', 'would', 'make', 'like',
            'into', 'him', 'has', 'more', 'go', 'no', 'way', 'could', 'my',
            'than', 'first', 'water', 'been', 'call', 'who', 'its', 'now',
            'find', 'long', 'down', 'day', 'did', 'get', 'come', 'made',
            'may', 'part', 'over', 'new', 'sound', 'take', 'only', 'little',
            'work', 'know', 'place', 'year', 'live', 'me', 'back', 'give',
            'most', 'very', 'after', 'thing', 'our', 'name', 'good', 'sentence',
            'man', 'think', 'say', 'great', 'where', 'help', 'through', 'much',
            'before', 'line', 'right', 'too', 'mean', 'old', 'any', 'same',
            'tell', 'boy', 'follow', 'came', 'want', 'show', 'also', 'around',
            'form', 'three', 'small', 'set', 'put', 'end', 'why', 'again',
            'turn', 'here', 'off', 'went', 'old', 'number', 'great', 'tell',
            'men', 'say', 'small', 'every', 'found', 'still', 'between',
            'mane', 'should', 'home', 'big', 'give', 'air', 'line', 'set',
            'own', 'under', 'read', 'last', 'never', 'us', 'left', 'end',
            'along', 'while', 'might', 'next', 'sound', 'below', 'saw',
            'something', 'thought', 'both', 'few', 'those', 'always', 'looked',
            'show', 'large', 'often', 'together', 'asked', 'house', 'world',
            'going', 'want', 'school', 'important', 'until', 'without', 'form',
            'black', 'white', 'words', 'students', 'during', 'started', 'include',
            'young', 'book', 'example', 'took', 'being', 'different', 'state',
            'never', 'became', 'between', 'high', 'really', 'something', 'most',
            'another', 'much', 'family', 'own', 'out', 'leave', 'put', 'old',
            'while', 'mean', 'on', 'keep', 'student', 'why', 'let', 'great',
            'same', 'big', 'group', 'begin', 'seem', 'country', 'help', 'talk',
            'where', 'turn', 'problem', 'every', 'start', 'hand', 'might',
            'american', 'show', 'part', 'about', 'against', 'place', 'over',
            'such', 'again', 'few', 'case', 'most', 'week', 'company', 'where',
            'system', 'each', 'right', 'program', 'hear', 'so', 'question',
            'during', 'work', 'play', 'government', 'run', 'small', 'number',
            'off', 'always', 'move', 'like', 'night', 'live', 'mr', 'point',
            'believe', 'hold', 'today', 'bring', 'happen', 'next', 'without',
            'before', 'large', 'all', 'million', 'must', 'home', 'under', 'water',
            'room', 'write', 'mother', 'area', 'national', 'money', 'story',
            'young', 'fact', 'month', 'different', 'lot', 'right', 'study',
            'book', 'eye', 'job', 'word', 'though', 'business', 'issue', 'side',
            'kind', 'four', 'head', 'far', 'black', 'long', 'both', 'little',
            'house', 'yes', 'after', 'since', 'long', 'provide', 'service',
            'around', 'friend', 'important', 'father', 'sit', 'away', 'until',
            'power', 'hour', 'game', 'often', 'yet', 'line', 'political', 'end',
            'among', 'ever', 'stand', 'bad', 'lose', 'however', 'member', 'pay',
            'law', 'meet', 'car', 'city', 'almost', 'include', 'continue',
            'set', 'later', 'community', 'much', 'name', 'five', 'once', 'white',
            'least', 'president', 'learn', 'real', 'change', 'team', 'minute',
            'best', 'several', 'idea', 'kid', 'body', 'information', 'back',
            'parent', 'face', 'others', 'level', 'office', 'door', 'health',
            'person', 'art', 'war', 'history', 'party', 'within', 'grow',
            'result', 'open', 'change', 'morning', 'walk', 'reason', 'low',
            'win', 'research', 'girl', 'guy', 'early', 'food', 'before', 'moment',
            'himself', 'air', 'teacher', 'force', 'offer'
        ]
        return set(basic_words)
    
    def analyze_text_difficulty(self, text: str) -> Dict[str, Any]:
        """分析文本难度"""
        # 简单的词汇分析
        words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
        total_words = len(words)
        unique_words = len(set(words))
        
        # 句子分析
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        # 计算常用词比例
        common_word_count = sum(1 for word in words if word in self.common_words)
        common_word_ratio = common_word_count / total_words if total_words > 0 else 0
        
        # 计算平均句长
        avg_sentence_length = total_words / len(sentences) if sentences else 0
        
        # 识别难词
        difficult_words = [word for word in set(words) 
                         if word not in self.common_words and len(word) > 3]
        
        # 计算难度评分 (1-10, 10最难)
        difficulty_score = self._calculate_difficulty_score(
            common_word_ratio, avg_sentence_length, len(difficult_words), unique_words
        )
        
        return {
            'total_words': total_words,
            'unique_words': unique_words,
            'common_word_ratio': common_word_ratio,
            'avg_sentence_length': avg_sentence_length,
            'difficult_words': difficult_words[:15],
            'difficulty_score': difficulty_score,
            'reading_level': self._get_reading_level(difficulty_score),
            'estimated_reading_time': self._estimate_reading_time(total_words),
            'vocabulary_coverage': common_word_ratio * 100
        }
    
    def _calculate_difficulty_score(self, common_ratio: float, avg_sent_len: float, 
                                  difficult_count: int, unique_count: int) -> float:
        """计算难度评分"""
        # 基于研究文献的98%词汇覆盖率原则
        coverage_penalty = max(0, (0.98 - common_ratio) * 8)
        sentence_penalty = max(0, (avg_sent_len - 15) * 0.15)
        difficulty_penalty = min(difficult_count * 0.08, 2.5)
        
        base_score = 5
        total_penalty = coverage_penalty + sentence_penalty + difficulty_penalty
        
        return min(10, max(1, base_score + total_penalty))
    
    def _get_reading_level(self, score: float) -> str:
        """根据评分获取阅读水平"""
        if score <= 3:
            return "初级 (适合初学者)"
        elif score <= 5:
            return "初中级 (适合有基础的学习者)"
        elif score <= 7:
            return "中级 (适合中等水平学习者)"
        elif score <= 8.5:
            return "中高级 (需要较好的英语基础)"
        else:
            return "高级 (需要扎实的英语功底)"
    
    def _estimate_reading_time(self, word_count: int) -> str:
        """估算阅读时间"""
        minutes = word_count / 130  # 中国EFL学习者平均阅读速度
        
        if minutes < 1:
            return f"{int(minutes * 60)}秒"
        elif minutes < 60:
            return f"{int(minutes)}分钟"
        else:
            hours = int(minutes / 60)
            remaining_minutes = int(minutes % 60)
            return f"{hours}小时{remaining_minutes}分钟"

class EnhancedNovelReader:
    """增强版小说阅读助手"""
    
    def __init__(self, model_name: str = "huihui_ai/qwenlong-abliterated:latest"):
        self.model_name = model_name
        self.ollama_url = "http://localhost:11434/api/generate"
        self.processed_paragraphs = []
        self.difficulty_analyzer = TextDifficultyAnalyzer()
        self.vocab_db = VocabularyDatabase()
        # 可用模型列表
        self.available_models = [
            "huihui_ai/qwenlong-abliterated:latest",
            "gemma3:12b",
            "gemma3:27b", 
            "qwen3:32b",
            "qwen3:8b",
            "gemma3:4b",
            "phi4:latest"
        ]
    
    def set_model(self, model_name: str):
        """设置使用的模型"""
        if model_name in self.available_models:
            self.model_name = model_name
            logger.info(f"模型已切换为: {model_name}")
        else:
            logger.warning(f"模型 {model_name} 不在可用列表中")
    
    def create_enhanced_analysis_prompt(self, paragraph: str, difficulty_info: Dict) -> str:
        """创建增强的分析提示词（用于单段落详细分析，保留思维链）"""
        prompt = f"""
作为英语教学专家，请对以下英文小说段落进行深度分析，特别关注中国学生的学习需求：

【原文段落】
{paragraph}

【段落基本信息】
- 总词数：{difficulty_info['total_words']}
- 独特词汇：{difficulty_info['unique_words']}
- 词汇覆盖率：{difficulty_info['vocabulary_coverage']:.1f}%
- 难度等级：{difficulty_info['reading_level']}
- 预估阅读时间：{difficulty_info['estimated_reading_time']}

请按照以下结构进行详细分析：

## 📊 文本难度评估
- 根据98%词汇覆盖率原则，评估此段落对中国学生的难度
- 指出可能造成理解障碍的语言特征

## 📚 核心词汇深度解析
请选择5-8个关键词汇进行深度分析，包括：
- 词汇的基本含义和词性
- 词族关系（如：inform → information, informative, informant）
- 常用搭配和固定用法
- 在不同语境中的含义变化
- 同义词和反义词

## 🏛️ 语言结构分析
- 识别复杂句式结构（如倒装、省略、嵌套等）
- 解释可能困扰中国学生的语法现象
- 对比中英语言差异

## 🌍 文化背景深度解读
- 详细解释西方文化背景知识
- 历史、社会、宗教背景说明
- 帮助理解文本中的文化内涵

## 🎭 文学技巧与小说要素
- 叙事视角和技巧分析
- 人物塑造和性格描写
- 情节发展和文学手法
- 象征意义和隐喻解读

## 🎯 阅读策略指导
基于研究建议，提供具体的阅读策略：
- 预测和推理技巧
- 上下文推断方法
- 如何处理生词
- 提高阅读流畅度的建议

## 💡 思考问题
设计3-5个深层思考问题，促进：
- 批判性思维
- 文本理解
- 跨文化比较
- 个人反思

## 🧠 记忆与理解检查
- 段落核心信息概括
- 关键细节回顾
- 理解程度自测问题

## 🈶 精准中文翻译
提供两个版本的翻译：
1. 直译版本（保持原文结构）
2. 意译版本（符合中文表达习惯）

请确保分析详细、准确，特别关注中国学生的学习特点和需求。
"""
        return prompt
    
    def create_simplified_analysis_prompt(self, paragraph: str, difficulty_info: Dict) -> str:
        """创建简化的分析提示词（用于整本书处理，关闭思维链以提高速度）"""
        prompt = f"""
请对以下英文小说段落进行快速分析，为中国学生提供关键信息：

【原文段落】
{paragraph}

【段落信息】词数：{difficulty_info['total_words']}，难度：{difficulty_info['reading_level']}

请提供简洁分析：

## 📚 关键词汇（3-5个）
选择最重要的词汇，简要说明含义和用法。

## 🌍 文化背景
简要说明重要的文化背景知识。

## 🎭 文学要素
简要分析叙事技巧和文学手法。

## 🈶 中文翻译
提供准确的中文翻译。

请保持简洁，重点突出核心内容。
"""
        return prompt
    
    def call_ollama(self, prompt: str, is_simplified: bool = False) -> str:
        """调用ollama模型
        
        Args:
            prompt: 提示词
            is_simplified: 是否为简化分析（用于优化参数）
        """
        try:
            # 根据分析类型调整参数
            if is_simplified:
                # 简化分析：更低的温度，更少的tokens，更短的超时
                options = {
                    "temperature": 0.1,
                    "top_p": 0.8,
                    "max_tokens": 2000,
                    "repeat_penalty": 1.0,
                }
                timeout = 120  # 更短的超时时间
            else:
                # 详细分析：标准参数
                options = {
                    "temperature": 0.3,
                    "top_p": 0.9,
                    "max_tokens": 6000,
                    "repeat_penalty": 1.1,
                }
                timeout = 300
            
            payload = {
                "model": self.model_name,
                "prompt": prompt,
                "stream": False,
                "options": options
            }
            
            response = requests.post(self.ollama_url, json=payload, timeout=timeout)
            
            if response.status_code == 200:
                result = response.json()
                return result.get('response', '')
            else:
                logger.error(f"Ollama API error: {response.status_code}")
                return f"错误：API调用失败，状态码：{response.status_code}"
                
        except Exception as e:
            logger.error(f"Error calling Ollama: {str(e)}")
            return f"错误：{str(e)}"
    
    def analyze_paragraph(self, paragraph: str, index: int, use_detailed_analysis: bool = True) -> Dict[str, Any]:
        """分析段落
        
        Args:
            paragraph: 要分析的段落文本
            index: 段落索引
            use_detailed_analysis: 是否使用详细分析（True=详细分析，False=简化分析）
        """
        analysis_type = "详细" if use_detailed_analysis else "简化"
        logger.info(f"正在进行{analysis_type}分析第 {index + 1} 段落...")
        
        # 进行难度分析
        difficulty_info = self.difficulty_analyzer.analyze_text_difficulty(paragraph)
        
        # 根据分析类型选择提示词
        if use_detailed_analysis:
            prompt = self.create_enhanced_analysis_prompt(paragraph, difficulty_info)
        else:
            prompt = self.create_simplified_analysis_prompt(paragraph, difficulty_info)
        
        # 获取AI分析
        analysis = self.call_ollama(prompt, is_simplified=not use_detailed_analysis)
        
        # 提取并保存词汇
        self._extract_and_save_vocabulary(paragraph, analysis)
        
        result = {
            "index": index + 1,
            "original_text": paragraph,
            "difficulty_info": difficulty_info,
            "analysis": analysis,
            "analysis_type": analysis_type,
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        self.processed_paragraphs.append(result)
        return result
    
    def _extract_and_save_vocabulary(self, text: str, analysis: str):
        """从文本和分析中提取词汇并保存到数据库"""
        # 简单的词汇提取（可以后续改进）
        words = word_tokenize(text.lower())
        words = [word for word in words if word.isalpha() and len(word) > 3]
        
        for word in set(words):
            if word not in self.difficulty_analyzer.common_words:
                # 这里可以添加更复杂的词汇定义提取逻辑
                self.vocab_db.add_word(word, "", "", 5)
    
    def get_reading_recommendations(self, difficulty_score: float) -> List[str]:
        """根据难度评分提供阅读建议"""
        recommendations = []
        
        if difficulty_score > 8:
            recommendations.extend([
                "🚨 此文本难度较高，建议：",
                "• 先预习关键词汇和文化背景",
                "• 分段阅读，不要急于求成",
                "• 使用词典和在线资源辅助理解",
                "• 考虑先阅读简化版本"
            ])
        elif difficulty_score > 6:
            recommendations.extend([
                "⚠️ 此文本具有一定挑战性，建议：",
                "• 在阅读前浏览一遍，了解大意",
                "• 重点关注段落主题句",
                "• 遇到生词先尝试猜测含义",
                "• 适当使用词典辅助"
            ])
        else:
            recommendations.extend([
                "✅ 此文本难度适中，建议：",
                "• 保持流畅阅读，不要过分纠结生词",
                "• 注意文章的逻辑结构",
                "• 尝试预测后续情节发展",
                "• 享受阅读过程"
            ])
        
        return recommendations
    
    def split_text_into_paragraphs(self, text: str) -> List[str]:
        """智能分割文本为段落"""
        paragraphs = re.split(r'\n\s*\n', text.strip())
        
        cleaned_paragraphs = []
        for p in paragraphs:
            p = p.strip()
            if p and len(p.split()) >= 25:  # 至少25个词
                cleaned_paragraphs.append(p)
        
        return cleaned_paragraphs
    
    def create_enhanced_docx(self, novel_title: str = "英文小说阅读分析报告") -> str:
        """创建增强版DOCX文档"""
        doc = Document()
        
        title = doc.add_heading('📚 ' + novel_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"生成时间：{time.strftime('%Y年%m月%d日 %H:%M:%S')}")
        doc.add_paragraph(f"处理段落数：{len(self.processed_paragraphs)}")
        doc.add_paragraph("基于《How to read English novels for Chinese students》研究文献")
        
        # 添加分析模式信息
        if self.processed_paragraphs:
            analysis_mode = self.processed_paragraphs[0].get('analysis_type', '详细')
            doc.add_paragraph(f"分析模式：{analysis_mode}分析")
            if analysis_mode == "简化":
                doc.add_paragraph("⚡ 采用快速分析模式，关闭思维链以提高处理速度")
        
        doc.add_paragraph("=" * 60)
        
        if self.processed_paragraphs:
            # 添加总体统计
            doc.add_heading('📊 阅读统计概览', level=1)
            
            total_words = sum(p['difficulty_info']['total_words'] for p in self.processed_paragraphs)
            avg_difficulty = sum(p['difficulty_info']['difficulty_score'] for p in self.processed_paragraphs) / len(self.processed_paragraphs)
            
            doc.add_paragraph(f"• 总词数：{total_words}")
            doc.add_paragraph(f"• 平均难度评分：{avg_difficulty:.1f}/10")
            doc.add_paragraph(f"• 预估总阅读时间：{self.difficulty_analyzer._estimate_reading_time(total_words)}")
            
            # 添加每个段落的分析
            for paragraph_data in self.processed_paragraphs:
                doc.add_page_break()
                
                doc.add_heading(f"第 {paragraph_data['index']} 段", level=1)
                
                difficulty_info = paragraph_data['difficulty_info']
                doc.add_heading('📊 难度评估', level=2)
                doc.add_paragraph(f"难度评分：{difficulty_info['difficulty_score']:.1f}/10")
                doc.add_paragraph(f"阅读等级：{difficulty_info['reading_level']}")
                doc.add_paragraph(f"词汇覆盖率：{difficulty_info['vocabulary_coverage']:.1f}%")
                
                doc.add_heading('📖 原文', level=2)
                p = doc.add_paragraph(paragraph_data['original_text'])
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                doc.add_heading('🔍 详细分析', level=2)
                analysis_p = doc.add_paragraph(paragraph_data['analysis'])
                analysis_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                recommendations = self.get_reading_recommendations(difficulty_info['difficulty_score'])
                if recommendations:
                    doc.add_heading('💡 阅读建议', level=2)
                    for rec in recommendations:
                        doc.add_paragraph(rec)
        
        filename = f"enhanced_novel_analysis_{time.strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        return filename

class EnhancedGradioInterface:
    """增强版Gradio界面"""
    
    def __init__(self):
        self.reader = EnhancedNovelReader()
        self.current_paragraphs = []
        self.current_index = 0
        self.current_novel_title = "未命名小说"
        self.current_model = self.reader.model_name
    
    def change_model(self, model_name: str) -> str:
        """切换模型"""
        self.reader.set_model(model_name)
        self.current_model = model_name
        return f"✅ 已切换到模型：{model_name}"
        
    def handle_file_upload(self, uploaded_file_path) -> Tuple[str, str]:
        """处理上传的文件"""
        try:
            if uploaded_file_path is None or uploaded_file_path == "":
                return "❌ 请选择要上传的文件", ""
            
            # 检查文件类型
            if not uploaded_file_path.lower().endswith(('.txt', '.md')):
                return "❌ 只支持 .txt 和 .md 格式的文件", ""
            
            # 获取文件名
            self.current_novel_title = os.path.splitext(os.path.basename(uploaded_file_path))[0]
            
            # 读取文件内容
            with open(uploaded_file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return self._load_content(content)
            
        except Exception as e:
            return f"❌ 上传文件时出错：{str(e)}", ""
    
    def load_and_analyze_novel(self, file_path: str) -> Tuple[str, str]:
        """加载并分析小说"""
        try:
            if file_path and os.path.exists(file_path):
                self.current_novel_title = os.path.splitext(os.path.basename(file_path))[0]
                
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                return self._load_content(content)
            else:
                return "❌ 文件路径无效或文件不存在", ""
        except Exception as e:
            return f"❌ 加载文件时出错：{str(e)}", ""
    
    def _load_content(self, content: str) -> Tuple[str, str]:
        """加载内容的共同逻辑"""
        self.current_paragraphs = self.reader.split_text_into_paragraphs(content)
        self.current_index = 0
        self.reader.processed_paragraphs = []
        
        overall_difficulty = self.reader.difficulty_analyzer.analyze_text_difficulty(content)
        
        status_message = f"✅ 成功加载小说《{self.current_novel_title}》，共 {len(self.current_paragraphs)} 段落"
        
        difficulty_summary = f"""📊 整体难度分析：
• 总词数：{overall_difficulty['total_words']:,}
• 独特词汇：{overall_difficulty['unique_words']:,}
• 词汇覆盖率：{overall_difficulty['vocabulary_coverage']:.1f}%
• 难度评分：{overall_difficulty['difficulty_score']:.1f}/10
• 阅读等级：{overall_difficulty['reading_level']}
• 预估阅读时间：{overall_difficulty['estimated_reading_time']}

💡 阅读建议：
{chr(10).join(self.reader.get_reading_recommendations(overall_difficulty['difficulty_score']))}"""
        
        return status_message, difficulty_summary
    
    def process_next_paragraph(self) -> Tuple[str, str, str, str]:
        """处理下一个段落"""
        if not self.current_paragraphs:
            return "❌ 请先加载小说文件", "", "", ""
        
        if self.current_index >= len(self.current_paragraphs):
            return "✅ 所有段落已处理完成", "", "", ""
        
        current_paragraph = self.current_paragraphs[self.current_index]
        result = self.reader.analyze_paragraph(current_paragraph, self.current_index)
        self.current_index += 1
        
        progress_info = f"已处理 {self.current_index}/{len(self.current_paragraphs)} 段落"
        
        difficulty_info = result['difficulty_info']
        difficulty_display = f"""📊 当前段落难度：
• 难度评分：{difficulty_info['difficulty_score']:.1f}/10
• 阅读等级：{difficulty_info['reading_level']}
• 词汇覆盖率：{difficulty_info['vocabulary_coverage']:.1f}%
• 预估阅读时间：{difficulty_info['estimated_reading_time']}
• 总词数：{difficulty_info['total_words']}，独特词汇：{difficulty_info['unique_words']}"""
        
        return progress_info, difficulty_display, result['original_text'], result['analysis']
    
    def process_entire_novel(self) -> str:
        """处理整本小说（使用简化分析模式，关闭思维链以提高速度）"""
        if not self.current_paragraphs:
            return "❌ 请先加载小说文件"
        
        try:
            total_paragraphs = len(self.current_paragraphs)
            logger.info(f"开始处理整本小说《{self.current_novel_title}》，共 {total_paragraphs} 段落")
            logger.info("📈 使用简化分析模式，关闭思维链以提高处理速度")
            
            # 重置处理状态
            self.reader.processed_paragraphs = []
            
            # 处理所有段落 - 使用简化分析模式（关闭思维链）
            for i, paragraph in enumerate(self.current_paragraphs):
                logger.info(f"正在快速处理第 {i+1}/{total_paragraphs} 段落")
                # 使用 use_detailed_analysis=False 来关闭思维链
                result = self.reader.analyze_paragraph(paragraph, i, use_detailed_analysis=False)
            
            # 保存完整分析
            filename = self.reader.create_enhanced_docx(self.current_novel_title)
            
            final_message = f"""✅ 整本小说快速处理完成！

📖 小说名称：《{self.current_novel_title}》
📊 处理统计：共处理 {total_paragraphs} 个段落
📄 分析报告：已保存为 {filename}
⚡ 分析模式：简化模式（关闭思维链以提高速度）

🎯 分析包含：
• 每段落的关键词汇分析
• 重要文化背景说明
• 基础文学技巧分析
• 精准的中文翻译

📚 您可以打开 {filename} 查看完整的分析报告！

💡 提示：如需详细分析，请使用"处理下一段"功能逐段分析。"""
            
            logger.info(f"整本小说快速处理完成，报告保存为：{filename}")
            return final_message
            
        except Exception as e:
            error_message = f"❌ 处理过程中出现错误：{str(e)}"
            logger.error(error_message)
            return error_message
    
    def save_enhanced_analysis(self) -> str:
        """保存增强分析结果"""
        if not self.reader.processed_paragraphs:
            return "❌ 没有已处理的段落可以保存"
        
        filename = self.reader.create_enhanced_docx(self.current_novel_title)
        return f"✅ 增强分析报告已保存为 {filename}，共包含 {len(self.reader.processed_paragraphs)} 个段落的详细分析"

def create_enhanced_interface():
    """创建增强版Gradio界面"""
    interface = EnhancedGradioInterface()
    
    with gr.Blocks(title="英文小说阅读辅助软件", theme=gr.themes.Soft()) as demo:
        gr.Markdown("# 📚 英文小说阅读辅助软件")
        gr.Markdown("Designed by Toby")
        
        with gr.Row():
            with gr.Column(scale=1):
                gr.Markdown("## 🤖 模型设置")
                model_dropdown = gr.Dropdown(
                    choices=interface.reader.available_models,
                    value=interface.current_model,
                    label="选择AI模型",
                    info="不同模型有不同的特点，可根据需要切换"
                )
                model_status = gr.Textbox(label="模型状态", interactive=False)
                
                gr.Markdown("## 📁 文件加载与难度分析")
                
                # 添加文件上传选项
                with gr.Tabs():
                    with gr.TabItem("📁 上传文件"):
                        file_upload = gr.File(
                            label="选择小说文件 (.txt 或 .md)",
                            file_types=['.txt', '.md'],
                            type="filepath"
                        )
                        upload_btn = gr.Button("📤 上传并分析", variant="primary")
                    
                    with gr.TabItem("📂 本地文件"):
                        file_input = gr.Textbox(
                            label="小说文件路径",
                            placeholder="输入小说文件的完整路径...",
                            value="The English Patient.txt"
                        )
                        load_btn = gr.Button("📊 加载并分析", variant="primary")
                
                load_status = gr.Textbox(label="加载状态", interactive=False)
                
                difficulty_analysis = gr.Textbox(
                    label="📊 难度分析报告", 
                    lines=12, 
                    interactive=False,
                    placeholder="文本难度分析将在这里显示..."
                )
                
                gr.Markdown("## 🎯 处理选项")
                gr.Markdown("**📝 处理模式说明：**")
                gr.Markdown("• **处理下一段**：详细分析模式，包含完整思维链，适合深度学习")
                gr.Markdown("• **处理整本小说**：快速分析模式，关闭思维链，大幅提高处理速度")
                
                progress_info = gr.Textbox(label="处理进度", interactive=False)
                
                with gr.Row():
                    next_btn = gr.Button("➡️ 处理下一段（详细模式）", variant="secondary")
                    process_all_btn = gr.Button("🚀 处理整本小说（快速模式）", variant="primary")
                
                with gr.Row():
                    save_btn = gr.Button("💾 保存当前分析", variant="secondary")
        
        with gr.Row():
            with gr.Column(scale=2):
                current_difficulty = gr.Textbox(
                    label="📊 当前段落难度信息",
                    lines=6,
                    interactive=False,
                    placeholder="段落难度信息将在这里显示..."
                )
                
                gr.Markdown("## 📖 英文原文")
                original_text = gr.Textbox(
                    label="原文段落",
                    lines=8,
                    interactive=False,
                    placeholder="英文原文将在这里显示..."
                )
                
                gr.Markdown("## 🔍 深度分析结果")
                analysis_result = gr.Textbox(
                    label="专业分析",
                    lines=25,
                    interactive=False,
                    placeholder="详细的专业分析结果将在这里显示..."
                )
        
        # 事件绑定
        model_dropdown.change(
            fn=interface.change_model,
            inputs=[model_dropdown],
            outputs=[model_status]
        )
        
        upload_btn.click(
            fn=interface.handle_file_upload,
            inputs=[file_upload],
            outputs=[load_status, difficulty_analysis]
        )
        
        load_btn.click(
            fn=interface.load_and_analyze_novel,
            inputs=[file_input],
            outputs=[load_status, difficulty_analysis]
        )
        
        next_btn.click(
            fn=interface.process_next_paragraph,
            inputs=[],
            outputs=[progress_info, current_difficulty, original_text, analysis_result]
        )
        
        process_all_btn.click(
            fn=interface.process_entire_novel,
            inputs=[],
            outputs=[progress_info]
        )
        
        save_btn.click(
            fn=interface.save_enhanced_analysis,
            inputs=[],
            outputs=[progress_info]
        )
    
    return demo

if __name__ == "__main__":
    demo = create_enhanced_interface()
    demo.launch(
        server_name="127.0.0.1",
        server_port=7862,
        share=False,
        show_error=True
    ) 